"""Markdown to DOCX converter using python-docx and markdown-it-py."""

import re
import sys
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_markdown(md_text: str) -> list[dict]:
    """Parse markdown text into a list of block elements."""
    blocks = []
    lines = md_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]

        if re.match(r"^#{1,6}\s", line):
            match = re.match(r"^(#{1,6})\s+(.*)", line)
            level = len(match.group(1))
            blocks.append({"type": "heading", "level": level, "text": match.group(2)})
            i += 1

        elif line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|", lines[i + 1]):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            header = [c.strip() for c in table_lines[0].strip("|").split("|")]
            rows = []
            for tl in table_lines[2:]:
                rows.append([c.strip() for c in tl.strip("|").split("|")])
            blocks.append({"type": "table", "header": header, "rows": rows})

        elif re.match(r"^\d+\.\s", line):
            items = []
            while i < len(lines) and re.match(r"^\d+\.\s", lines[i]):
                items.append(re.sub(r"^\d+\.\s+", "", lines[i]))
                i += 1
            blocks.append({"type": "ordered_list", "items": items})

        elif line.startswith("- "):
            items = []
            while i < len(lines) and lines[i].startswith("- "):
                items.append(lines[i][2:])
                i += 1
            blocks.append({"type": "unordered_list", "items": items})

        elif line.strip() == "":
            i += 1

        else:
            para_lines = []
            while i < len(lines) and lines[i].strip() != "" and not lines[i].startswith("#") and not lines[i].startswith("|") and not lines[i].startswith("- ") and not re.match(r"^\d+\.\s", lines[i]):
                para_lines.append(lines[i])
                i += 1
            blocks.append({"type": "paragraph", "text": " ".join(para_lines)})

    return blocks


def add_rich_text(paragraph, text: str):
    """Add text with bold/italic markdown formatting to a paragraph."""
    pattern = re.compile(r"\*\*(.+?)\*\*|__(.+?)__|\*(.+?)\*|_(.+?)_|`(.+?)`")
    last_end = 0

    for match in pattern.finditer(text):
        if match.start() > last_end:
            paragraph.add_run(text[last_end:match.start()])

        if match.group(1) or match.group(2):
            run = paragraph.add_run(match.group(1) or match.group(2))
            run.bold = True
        elif match.group(3) or match.group(4):
            run = paragraph.add_run(match.group(3) or match.group(4))
            run.italic = True
        elif match.group(5):
            run = paragraph.add_run(match.group(5))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)

        last_end = match.end()

    if last_end < len(text):
        paragraph.add_run(text[last_end:])


def set_cell_shading(cell, color: str):
    """Set background color for a table cell."""
    from lxml import etree
    tc_pr = cell._tc.get_or_add_tcPr()
    shading_xml = (
        f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
        f' w:fill="{color}" w:val="clear"/>'
    )
    tc_pr.append(etree.fromstring(shading_xml))


def set_table_style(table):
    """Apply borders to a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    if tbl_pr is None:
        tbl_pr = ET.SubElement(tbl, qn("w:tblPr"))

    borders_xml = (
        '<w:tblBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    from lxml import etree
    borders_el = etree.fromstring(borders_xml)
    tbl_pr.append(borders_el)


def md_to_docx(md_path: str, docx_path: str = None):
    """Convert a markdown file to a DOCX document."""
    md_path = Path(md_path)
    if docx_path is None:
        docx_path = md_path.with_suffix(".docx")
    else:
        docx_path = Path(docx_path)

    md_text = md_path.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    heading_sizes = {1: 22, 2: 16, 3: 14, 4: 12, 5: 11, 6: 10.5}

    for block in blocks:
        if block["type"] == "heading":
            level = block["level"]
            p = doc.add_heading(level=level)
            run = p.add_run(block["text"])
            run.font.size = Pt(heading_sizes.get(level, 10.5))
            run.font.name = "Microsoft YaHei"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

        elif block["type"] == "paragraph":
            p = doc.add_paragraph()
            add_rich_text(p, block["text"])

        elif block["type"] == "ordered_list":
            for idx, item in enumerate(block["items"], 1):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.25)
                p.add_run(f"{idx}. ")
                add_rich_text(p, item)

        elif block["type"] == "unordered_list":
            for item in block["items"]:
                p = doc.add_paragraph(style="List Bullet")
                add_rich_text(p, item)

        elif block["type"] == "table":
            header = block["header"]
            rows = block["rows"]
            num_cols = len(header)
            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_style(table)

            for ci, h in enumerate(header):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h)
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_cell_shading(cell, "E8EDF2")

            for ri, row in enumerate(rows):
                for ci in range(min(len(row), num_cols)):
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(row[ci])
                    run.font.size = Pt(9)
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    doc.save(str(docx_path))
    print(f"Done: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python md_to_docx.py <input.md> [output.docx]")
        sys.exit(1)
    input_md = sys.argv[1]
    output_docx = sys.argv[2] if len(sys.argv) > 2 else None
    md_to_docx(input_md, output_docx)
